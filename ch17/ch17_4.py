# ch17_4.py
import docx

wdoc = docx.Document()
wdoc.add_heading('明志科大')            # 預設是標題1格式
wdoc.add_heading('明志科大', level=2)   # 標題2格式
wdoc.add_heading('明志科大', level=3)   # 標題3格式
print(wdoc.paragraphs[0].text)
print(wdoc.paragraphs[1].text)
print(wdoc.paragraphs[2].text)
wdoc.save('out17_4.docx')










