# ch17_1.py
import docx

wdoc = docx.Document('data17_1.docx')
print("段落數, 也可稱Paragraph物件數量 = ", len(wdoc.paragraphs))
for i in range(0, len(wdoc.paragraphs)):
    print("paragraph %d = " % i, wdoc.paragraphs[i].text)

print("Paragraph 1的Run數量 = ", len(wdoc.paragraphs[1].runs))
for i in range(0, len(wdoc.paragraphs[1].runs)):
    print("Run %d = " % i, wdoc.paragraphs[1].runs[i].text)

print("Paragraph 2的Run數量 = ", len(wdoc.paragraphs[2].runs))
for i in range(0, len(wdoc.paragraphs[2].runs)):
    print("Run %d = " % i, wdoc.paragraphs[2].runs[i].text)



