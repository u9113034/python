# ch17_5.py
import docx

wdoc = docx.Document()
ptr = wdoc.add_paragraph('我是第1段落')                     # 回傳ptr段落物件
print(wdoc.paragraphs[0].text)
wdoc.add_paragraph('我是第2段落')                           # 不加回傳也可以
print(wdoc.paragraphs[1].text)
prior_ptr = ptr.insert_paragraph_before('我是新的第1段落')  # 新段落插在ptr前面
print(wdoc.paragraphs[0].text)
wdoc.save('out17_5.docx')










