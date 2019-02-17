# ch17_7.py
import docx

wdoc = docx.Document()
ptr = wdoc.add_paragraph('我是第1段落')     # 回傳ptr段落物件
wdoc.add_paragraph('我是第2段落')           # 不加回傳也可以
wdoc.add_page_break()                       # 插入換頁      
wdoc.add_paragraph('我是新的頁開始段落')    # 新段落插在新的一頁
wdoc.save('out17_7.docx')










