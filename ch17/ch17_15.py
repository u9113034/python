# ch17_15.py
import docx

wdoc = docx.Document()
ptr = wdoc.add_paragraph('我是第1段落')  # 回傳ptr段落物件
run1 = ptr.add_run('我是粗體')           # 將run加入ptr段落物件
run1.bold = True                         # 設定粗體
run2 = ptr.add_run('我是斜體')           # 將run加入ptr段落物件
run2.italic = True                       # 設定斜體
wdoc.save('out17_15.docx')  










         
