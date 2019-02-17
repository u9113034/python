# ch17_8.py
import docx
from docx.shared import Inches,Cm

wdoc = docx.Document()
wdoc.add_paragraph('洪錦魁相片')     
wdoc.add_picture('洪錦魁.jpg', width=Inches(1.0), height=Cm(5.0))
wdoc.save('out17_8.docx')










