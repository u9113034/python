# ch17_14.py
import docx

wdoc = docx.Document()
wdoc.add_paragraph('明志科大', style='ListNumber')    # ListNumber
wdoc.add_paragraph('長庚科大', style='ListNumber')    # ListNumber
wdoc.add_paragraph('長庚大學', style='ListNumber')    # ListNumber
wdoc.add_paragraph('明志科大', style='ListBullet')    # ListBullet
wdoc.add_paragraph('長庚科大', style='ListBullet')    # ListBullet
wdoc.add_paragraph('長庚大學', style='ListBullet')    # ListBullet
wdoc.save('out17_14.docx')













