# ch17_6.py
import docx

wdoc = docx.Document()
ptr = wdoc.add_paragraph('我是第1段落') # 回傳ptr段落物件
print(wdoc.paragraphs[0].text)          # 列印段落
ptr.add_run('第一個run內容')            # 將run加入ptr段落物件
ptr.add_run('第二個新run內容')          # 將run加入ptr段落物件
print(wdoc.paragraphs[0].text)          # 列印段落
wdoc.save('out17_6.docx')










