# ex3.py 設計自傳
import docx

wdoc=docx.Document()
wdoc.add_heading('我的自傳')
wdoc.add_heading('姓名 周侊懋', level=2)
wdoc.add_heading('生日 1983/03/27', level=3)

wdoc.add_paragraph('學習程式如下表')
table=wdoc.add_table(rows=3,cols=3)
row=table.rows[0]
row.cells[0].text='Visual Basic'
row.cells[1].text='JAVA'
row.cells[2].text='C/C++'
row=table.rows[1]
row.cells[0].text='Python'
row.cells[1].text='HTML'
row.cells[2].text='Javascript'
row=table.rows[2]
row.cells[0].text='PHP'
row.cells[1].text='Android'
row.cells[2].text='Inventor2'
table.style='LightShading-Accent1'
wdoc.add_paragraph()
wdoc.add_paragraph('畢業學校如下')
wdoc.add_paragraph('國立雲林科技大學', style='ListBullet')
wdoc.add_paragraph('國立台南高工', style='ListBullet')
wdoc.add_paragraph('文賢國中', style='ListBullet')

wdoc.save('ex3_1.docx')