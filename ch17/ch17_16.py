# ch17_16.py
import docx
import re

wdoc = docx.Document('data17_16.docx')
txt = docx.Document()
pattern = r'CIA (\w)\w*'                # 欲搜尋FBI + 空一格後的名字        
newstr = r'\1***'                       # 新字串使用隱藏文字
txt.add_paragraph(re.sub(pattern,newstr,wdoc.paragraphs[0].text))
txt.save('out17_16.docx')







    



        


